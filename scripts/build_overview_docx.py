"""Generate docs/overview.docx for the manager-facing handoff.

Renders a clean PNG of the system architecture with matplotlib, then
builds a Word document with the same content as docs/overview.md plus
that embedded diagram. Run with:

    .venv/bin/python scripts/build_overview_docx.py
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.patches as patches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
DIAGRAM_PNG = REPO / "docs" / "_architecture-diagram.png"
OUT_DOCX = REPO / "docs" / "overview.docx"


# =============================================================================
# 1. Render the architecture diagram as a PNG
# =============================================================================
def _draw_box(ax, x, y, w, h, label, fill, *, edge="#0b1733", text="#0b1733",
              fontsize=9, weight="normal"):
    rect = patches.FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.06",
        linewidth=1.2, edgecolor=edge, facecolor=fill,
    )
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h / 2, label,
            ha="center", va="center",
            fontsize=fontsize, color=text, fontweight=weight,
            wrap=True)


def _arrow(ax, x1, y1, x2, y2, *, color="#395270", style="->", lw=1.2):
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle=style, color=color, lw=lw),
    )


def _draw_band(ax, x, y, w, h, label, fill, *, edge="#142647", text="#142647",
               fontsize=10):
    """Draw an outer band with the title in a header strip at the top
    so it never overlaps the inner boxes that get drawn inside it."""
    rect = patches.FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.06",
        linewidth=1.2, edgecolor=edge, facecolor=fill,
    )
    ax.add_patch(rect)
    # Place the label near the top edge of the band, not centered
    ax.text(x + w / 2, y + h - 0.32, label,
            ha="center", va="center",
            fontsize=fontsize, color=text, fontweight="bold")


def render_diagram(out: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 9), dpi=160)
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 13.5)
    ax.axis("off")

    # Color palette — soft, document-friendly
    NAVY = "#142647"
    NAVY_FILL = "#e7edf6"
    CYAN_FILL = "#dff7f3"
    AMBER_FILL = "#fbeed1"
    EMERALD_FILL = "#dff1e7"
    WHITE = "#ffffff"
    GREY_FILL = "#f3f5f9"

    # Title
    ax.text(8, 13.05, "HITL agent runtime — system architecture",
            ha="center", fontsize=13, fontweight="bold", color=NAVY)

    # =========== FRONTEND BAND (top) ===========
    _draw_band(ax, 0.4, 10.4, 15.2, 2.1, "Frontend (React + Vite)",
               fill=GREY_FILL)
    _draw_box(ax, 0.7, 10.55, 3.4, 1.25,
              "Operator console\n(prompt · chips · history)",
              fill=WHITE, fontsize=8)
    _draw_box(ax, 4.3, 10.55, 3.4, 1.25,
              "FlowStage\n(4-stage envelope)",
              fill=WHITE, fontsize=8)
    _draw_box(ax, 7.9, 10.55, 3.4, 1.25,
              "Lineage panel\n(append-only events)",
              fill=WHITE, fontsize=8)
    _draw_box(ax, 11.5, 10.55, 3.9, 1.25,
              "Knowledge tile\n(Data sources · Ontologies · Actions)",
              fill=WHITE, fontsize=8)

    # arrow frontend → orchestrator
    _arrow(ax, 8, 10.35, 8, 9.7, lw=1.6)
    ax.text(8.2, 10.0, "REST + SSE  (JWT auth)", fontsize=8, color=NAVY,
            ha="left", va="center")

    # =========== ORCHESTRATOR ROW ===========
    _draw_box(ax, 0.4, 8.3, 15.2, 1.4,
              "Orchestrator — runs each case through the four stages\n"
              "agent intake → proposal → review → execute",
              fill=NAVY_FILL, fontsize=10, weight="bold")

    # =========== AGENT RUNTIME + 3 REGISTRIES ===========
    _draw_box(ax, 0.4, 5.7, 3.7, 2.0,
              "Agent runtime\n\n"
              "interpret_prompt\n"
              "(LLM classifier +\n"
              "NL fallback chain)",
              fill=AMBER_FILL, fontsize=8)

    _draw_box(ax, 4.3, 5.7, 3.7, 2.0,
              "Scenarios\n\n"
              "YAML recipes:\n"
              "facts + ontology_queries\n"
              "chips on the console",
              fill=CYAN_FILL, fontsize=8)

    _draw_box(ax, 8.2, 5.7, 3.7, 2.0,
              "Ontology + mappings\n\n"
              "business classes →\n"
              "source bindings;\n"
              "schema-aware NL parser",
              fill=CYAN_FILL, fontsize=8)

    _draw_box(ax, 12.1, 5.7, 3.5, 2.0,
              "Action registry\n\n"
              "typed write actions +\n"
              "executors\n"
              "(sql_update, http_request)",
              fill=CYAN_FILL, fontsize=8)

    # arrows from orchestrator down to each registry
    for cx in [2.25, 6.15, 10.05, 13.85]:
        _arrow(ax, cx, 8.25, cx, 7.75, lw=1.0)

    # =========== KNOWLEDGE RESOLVER LAYER ===========
    _draw_band(ax, 0.4, 3.0, 15.2, 2.0,
               "KnowledgeResolver Protocol — pluggable connector layer",
               fill=NAVY_FILL)
    names = ["CSV", "SQLite", "Postgres", "HTTP", "Vector store", "Neo4j"]
    cx_start = 0.65
    cx_w = 2.4
    cx_gap = 0.1
    for i, name in enumerate(names):
        x = cx_start + i * (cx_w + cx_gap)
        _draw_box(ax, x, 3.15, cx_w, 1.05, name, fill=WHITE, fontsize=9)

    # arrows registries → resolver layer
    _arrow(ax, 6.15, 5.65, 6.15, 5.05, lw=1.0)
    _arrow(ax, 10.05, 5.65, 10.05, 5.05, lw=1.0)
    _arrow(ax, 13.85, 5.65, 13.85, 5.05, lw=1.0)

    # =========== STORAGE / EXTERNAL ===========
    _draw_box(ax, 0.4, 1.0, 4.8, 1.7,
              "App SQLite store\n\n"
              "cases · append-only lineage · users",
              fill=EMERALD_FILL, fontsize=8)
    _draw_box(ax, 5.4, 1.0, 5.4, 1.7,
              "External enterprise sources\n\n"
              "CSVs · governance DBs ·\n"
              "vector embeddings · REST APIs",
              fill=EMERALD_FILL, fontsize=8)
    _draw_box(ax, 11.0, 1.0, 4.6, 1.7,
              "Neo4j (Bolt)\n\n"
              "all reads guarded by\n"
              "assert_read_only()",
              fill=EMERALD_FILL, fontsize=8)

    # arrows resolver → storage
    _arrow(ax, 3.0, 3.0, 3.0, 2.7, lw=0.9, color="#777")
    _arrow(ax, 8.0, 3.0, 8.0, 2.7, lw=0.9, color="#777")
    _arrow(ax, 13.5, 3.0, 13.5, 2.7, lw=0.9, color="#777")

    # Caption strip
    ax.text(0.4, 0.4,
            "Each layer is a Protocol/seam — connectors, executors, transports, "
            "and reviewer surfaces are all swappable.",
            fontsize=8, color="#555", style="italic")

    plt.tight_layout()
    plt.savefig(out, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# =============================================================================
# 2. Build the Word document
# =============================================================================
NAVY = RGBColor(0x14, 0x26, 0x47)
CYAN = RGBColor(0x00, 0x80, 0x80)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(8)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _para(doc: Document, text: str, *, bold_first_word: bool = False) -> None:
    p = doc.add_paragraph()
    if bold_first_word and " " in text:
        first, rest = text.split(" ", 1)
        run = p.add_run(first + " ")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        # Bold the leading "Term —" pattern when present
        if " — " in it:
            term, rest = it.split(" — ", 1)
            r1 = p.add_run(term)
            r1.bold = True
            p.add_run(" — " + rest)
        else:
            p.add_run(it)


def _table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # tighten font
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def build_doc(out: Path, diagram_png: Path) -> None:
    doc = Document()

    # Page-level defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    _h1(doc, "HITL Agent Runtime — System Overview")
    sub = doc.add_paragraph()
    r = sub.add_run("A non-technical walkthrough for stakeholders. "
                    "Pairs with docs/architecture.md (engineering view).")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    # ---- What this is ----
    _h2(doc, "What this is, in one paragraph")
    _para(doc,
          "A Human-in-the-Loop (HITL) agent runtime for enterprise workflows. "
          "Operators ask the agent to do work in plain English. The agent "
          "classifies the request, gathers the right data from connected "
          "systems, and either executes autonomously (read-only lookups, "
          "parameter changes within a pre-approved envelope) or pauses for a "
          "named human reviewer (anything irreversible — sanctions overrides, "
          "vendor onboarding, regulatory actions). Every step is recorded in "
          "an append-only audit trail.")

    # ---- Why it exists ----
    _h2(doc, "Why it exists")
    _para(doc, "Off-the-shelf agent platforms have one of two failure modes:")
    _bullets(doc, [
        "Too autonomous — they execute, but you can't see what knowledge informed each decision, and you can't insert a human approval step where you need one.",
        "Too rigid — they stop at \"auto-suggest the next form to fill in\" and don't actually take action.",
    ])
    _para(doc, "This system aims for the middle: the agent does real work, "
               "with real provenance, and with HITL review wired into the spine "
               "of every risky action.")

    # ---- Architecture diagram ----
    _h2(doc, "Architecture at a glance")
    _para(doc, "The diagram below shows the full system. Each row is a layer, "
               "and arrows show how a request flows from operator (top) down "
               "through the connected data systems (bottom).")
    doc.add_picture(str(diagram_png), width=Inches(6.5))
    cap = doc.add_paragraph()
    r = cap.add_run("Figure 1. Layered architecture. The four cyan/amber boxes "
                    "in the middle are the operator's mental model: scenarios, "
                    "ontology, actions, and the agent runtime that ties them "
                    "together.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Three audiences ----
    _h2(doc, "Three audiences and what each gets")
    _table(doc,
        ["Audience", "What they touch", "What they get"],
        [
            ["Operator",
             "Chat composer + a row of suggested-prompt chips",
             "A natural-language way to kick off pre-approved workflows + ad-hoc data queries against any registered source. No SQL or ontology authoring required."],
            ["Reviewer",
             "A Teams Adaptive Card (or in-app card) per pending decision",
             "Approve / Reject / Need more info, with the full context (active policy, actor scope, master data, prior similar cases, SOPs) bound into the card automatically."],
            ["Engineer / scenario author",
             "YAML files + a Knowledge admin tile",
             "Add a new connected system and the agent can use it across every existing scenario without rewriting any of them. Author one ontology, get many scenarios for free."],
        ],
    )

    # ---- Mental model ----
    _h2(doc, "The mental model")
    _para(doc, "Three nouns + three verbs.")
    _h3(doc, "Nouns (what the system knows about)")
    _bullets(doc, [
        "Data sources — connections to things that hold data. CSVs, SQLite, Postgres, HTTP APIs, vector stores, Neo4j graphs.",
        "Ontology — a business-language description of the entities the operator cares about (Supplier, Product, PurchaseOrder, etc.) and how each entity maps to columns in the data sources. Authored once; reused everywhere.",
        "Scenarios — recipes the agent follows. Each one is a pattern: the operator typed something like X → here's what the agent should fetch, what it should propose, who reviews it, and what happens on approve.",
    ])
    _h3(doc, "Verbs (what happens to a request)")
    _bullets(doc, [
        "Classify — the agent's LLM picks the scenario that best matches the operator's prompt.",
        "Bind — the framework gathers the knowledge each stage of the scenario needs (active policy, actor's scope, master-data record, prior decisions, SOPs) into a typed envelope. Every fact carries provenance.",
        "Decide — either auto-execute (low-risk, deterministic) or pause for a named human reviewer.",
    ])
    _para(doc, "That's it. Everything else is implementation.")

    # ---- Walkthrough ----
    _h2(doc, "A walkthrough: one operator request")
    quote = doc.add_paragraph()
    r = quote.add_run('"Override the SC-TC-001 sanctions block on order '
                      'ORD-44216 — customer says they have an OFAC license."')
    r.italic = True
    r.font.color.rgb = CYAN
    quote.paragraph_format.left_indent = Inches(0.4)

    steps = [
        ("Step 1 — Operator types it.",
         "The chat composer sends the prompt to the backend."),
        ("Step 2 — The agent classifies.",
         "The LLM looks at the catalog of scenarios, picks "
         "SC-TC-007 (Trade override — sanctioned counterparty), and "
         "rephrases the request. The operator sees the rephrasing as a "
         "clarifier and clicks Confirm."),
        ("Step 3 — The framework binds knowledge.",
         "Two stages run in sequence: the agent intake stage pulls the "
         "active TC-override policy + the agent's IAM scopes; the proposal "
         "stage pulls the product master record (encryption module P-EL-9001), "
         "the OFAC sanctions hit on the counterparty, and the contract value."),
        ("Step 4 — Policy decides.",
         "Sanctions overrides require named compliance officer approval, so "
         "the case enters the review stage. A Teams Adaptive Card is rendered "
         "with all the bound facts plus prior similar overrides and three "
         "relevant SOP excerpts retrieved by semantic search."),
        ("Step 5 — Reviewer decides.",
         "The named compliance officer reads the card. If they approve, the "
         "agent proceeds with the override and the case completes. If they "
         "reject or ask for more info, the case captures the rationale."),
        ("Step 6 — Audit trail.",
         "Every fact bound, every query issued, every decision is appended "
         "to an immutable log keyed by case id. Replayable later for "
         "compliance review."),
    ]
    for title, body in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(title)
        r.bold = True
        p.add_run(" " + body)

    _para(doc, "The whole loop happens in ~10 seconds for autonomous cases, "
               "or pauses indefinitely at the review step for HITL cases. The "
               "reviewer's decision arrives through a webhook (or the in-app "
               "surface today) and the orchestrator picks the case back up "
               "where it left off.")

    # ---- What's been built ----
    _h2(doc, "What's been built")
    _h3(doc, "Operator-facing")
    _bullets(doc, [
        "Suggested-prompt chips sorted newest-first so freshly added recipes surface to the top.",
        "Two opt-in fallback toggles below the composer: Try ontology lookup if no scenario matches (one-shot read-only data lookup) and Try write action (HITL-gated NL writes).",
        "Live envelope view — the operator can watch each stage's bound facts populate as the agent works.",
        "History panel — past conversations, search, replay with a forced reviewer decision.",
    ])
    _h3(doc, "Knowledge admin (the \"Knowledge\" tile)")
    _bullets(doc, [
        "Data sources tab — register a CSV, SQLite, Postgres, HTTP API, vector store, or Neo4j graph. Test the connection. Run ad-hoc SQL or Cypher in a playground.",
        "Ontologies tab — upload a YAML/JSON ontology defining business entities. Click \"Suggest mappings\" to have the LLM propose how each business attribute maps to columns in registered sources, then review and confirm. Per-class \"Generate lookup chip\" button creates an autonomous chip on the operator console.",
        "Actions tab — view registered write actions, preview what the LLM would extract from a natural-language prompt, remove non-default ones.",
    ])
    _h3(doc, "Audit and observability")
    _bullets(doc, [
        "Append-only lineage per case: every fact fetched, every query issued, every decision. Survives restart.",
        "Per-fact provenance in the reviewer card: each fact shows both what was asked (the ontology class) and which source answered (the physical data source).",
        "Live metrics dashboard — totals, decisions-by-scenario, top rejection reasons, cases-per-day sparkline.",
        "CSV export of cases + lineage with date-range filters.",
    ])
    _h3(doc, "Safety")
    _bullets(doc, [
        "HITL by default for write actions. The natural-language write path always sends the proposed action to a human reviewer; the executor only runs after approve.",
        "Read-only Cypher guard on every Neo4j query path. Writes via Cypher are rejected at the application layer.",
        "Default-flagged knowledge artefacts (built-in ontologies, seeded actions) refuse delete via the API.",
        "JWT auth + bcrypt passwords + token blocklist on logout. Per-user case isolation; admin role bypasses scoping.",
        "Rate limiting on login + register (10/min/IP) + app-wide 120/min default.",
    ])

    # ---- Differentiation ----
    _h2(doc, "Why it's different from a generic agent platform")
    _table(doc,
        ["Generic agent platform", "This system"],
        [
            ["LLM picks a function, calls it, returns.",
             "LLM picks a scenario; the framework binds the right knowledge at each stage and routes through HITL when policy says to."],
            ["Audit trail is a chat transcript.",
             "Audit trail is a typed envelope: every fact has provenance, every decision has rationale, every step is replayable."],
            ["Adding a new data source means re-prompting the LLM with new tool definitions.",
             "Adding a new data source means registering a connection and clicking \"Suggest mappings.\" Existing scenarios pick it up via the ontology layer without changes."],
            ["Human review is a chat message you can ignore.",
             "Human review is a structured Teams Adaptive Card with three buttons; the case literally pauses until a named reviewer decides."],
            ["Write actions go through the LLM.",
             "Write actions are operator-authored YAML with typed argument schemas and named executors. The LLM only fills the arguments; the executor is hand-written."],
        ],
    )

    # ---- Quality gates ----
    _h2(doc, "Quality gates today")
    _bullets(doc, [
        "214 automated tests passing (124 backend + 90 framework).",
        "Frontend type-check clean and production bundle builds.",
        "Live demo verified through every chip + the Neo4j graph + the ontology Query playground + the NL write-action preview.",
    ])

    # ---- Recommended next ----
    _h2(doc, "Recommended next steps")
    _h3(doc, "Implementable in 2 weeks (pilot prep + obvious feature gaps)")
    _bullets(doc, [
        "Pick one real customer data source (e.g. their Postgres governance store) and run the full flow: register → AI-map → review with a domain expert → confirm → click \"Generate lookup chip\" → demo. Proves the value prop on real data.",
        "Tighten the production deployment per docs/production.md — generate JWT secret, rotate the seeded admin password, set up HTTPS, configure backups.",
        "Decide the mapping governance model — who can confirm a mapping? Mappings are effectively access-control documents; treat the confirm step like an ACL change.",
        "RDF/Turtle ontology import (~2 days). For teams that already maintain ontologies in industry-standard formats.",
        "Bulk \"generate all chips\" in the Mappings tab — today operators click per class. Quality-of-life improvement.",
        "Cypher action executor (~2 days). The action registry has SQL + HTTP today. Adding a Cypher write executor lets operators kick off graph mutations through the same HITL path.",
    ])
    _h3(doc, "Production readiness (one quarter)")
    _bullets(doc, [
        "Move from in-memory to durable transport. Single worker with in-memory queues today; production scale needs Redis (case store) and Kafka/SQS (review-decision queue). The framework was designed for this swap; ~2 weeks.",
        "Observability — structured JSON logging, Prometheus metrics, audit-log retention policy. ~1 week.",
        "Cross-source entity resolution. Currently if Supplier S-001 lives in two sources, both records show up separately. Adding ER rules would dedupe with confidence scores. Defer until a customer asks; ~2 weeks of real engineering work.",
    ])
    _para(doc, "The architecture is intentionally built so each of these is "
               "additive, not a rewrite — the resolver layer, the binder "
               "protocols, and the mapping/ontology contracts are all "
               "\"swap-points\" the framework was designed around.")

    # ---- Save ----
    doc.save(str(out))


# =============================================================================
# Entry point
# =============================================================================
if __name__ == "__main__":
    print(f"Rendering diagram → {DIAGRAM_PNG}")
    render_diagram(DIAGRAM_PNG)
    print(f"Building Word doc → {OUT_DOCX}")
    build_doc(OUT_DOCX, DIAGRAM_PNG)
    print(f"Done. {OUT_DOCX.stat().st_size // 1024} KB")
